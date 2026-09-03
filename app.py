import io
import json
import os
import re
from typing import Any, Dict, List

import streamlit as st
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from google import genai
from google.genai import types

load_dotenv()

st.set_page_config(
    page_title="ResumeMatch AI",
    page_icon="📄",
    layout="wide",
)

MAX_TEXT_CHARS = 45000
MODEL = "gemini-2.5-flash"


def extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".docx"):
        return extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def clean_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def build_prompt(resume: str, jd: str, task: str) -> str:
    base_rules = """
You are an expert ATS resume analyzer, technical recruiter, and professional resume writer.

Analyze the candidate resume against the supplied job description.

NON-NEGOTIABLE RULES:
1. Never invent, assume, or fabricate candidate experience, skills, metrics, employers, dates, education, certifications, or achievements.
2. A missing keyword means it was not demonstrated in the supplied resume; it does NOT prove the candidate lacks the skill.
3. Never keyword-stuff a resume. Recommendations must be natural and relevant.
4. Do not change factual information such as company names, job titles, dates, degree names, or certifications.
5. When tailoring the resume, only use facts supported by the original resume.
6. Do not create fake metrics. If a useful metric is missing, suggest that the candidate add one if they can truthfully provide it.
7. ATS score is an estimate of compatibility, not a guarantee for any particular ATS vendor.
8. Prefer evidence and context over simple word frequency.
"""

    if task == "analyze":
        return f"""
{base_rules}

SCORING:
- Keyword Match: 25
- Required Skills Match: 20
- Responsibilities Match: 20
- Experience Match: 15
- Education/Certification Match: 10
- ATS Formatting/Parseability: 10
Total: 100.

Return ONLY valid JSON matching this schema:
{{
  "overall_score": 0,
  "match_level": "Strong Match | Moderate Match | Low Match",
  "score_breakdown": {{
    "keyword_match": {{"score": 0, "max_score": 25, "explanation": ""}},
    "required_skills": {{"score": 0, "max_score": 20, "explanation": ""}},
    "responsibilities_match": {{"score": 0, "max_score": 20, "explanation": ""}},
    "experience_match": {{"score": 0, "max_score": 15, "explanation": ""}},
    "education_certification": {{"score": 0, "max_score": 10, "explanation": ""}},
    "ats_formatting": {{"score": 0, "max_score": 10, "explanation": ""}}
  }},
  "job_requirements": {{
    "required_skills": [],
    "preferred_skills": [],
    "soft_skills": [],
    "experience_requirements": [],
    "education_requirements": [],
    "certifications": [],
    "responsibilities": [],
    "tools_and_technologies": []
  }},
  "keyword_analysis": {{
    "matching_keywords": [],
    "partial_keywords": [],
    "missing_keywords": []
  }},
  "strengths": [],
  "gaps": [],
  "formatting_issues": [],
  "recommendations": [
    {{
      "priority": "CRITICAL | HIGH | MEDIUM | LOW",
      "problem": "",
      "why_it_matters": "",
      "recommendation": "",
      "example": ""
    }}
  ],
  "bullet_improvements": [
    {{"original": "", "improved": "", "reason": ""}}
  ],
  "top_5_improvements": [],
  "final_recommendation": ""
}}

RESUME:
{resume[:MAX_TEXT_CHARS]}

JOB DESCRIPTION:
{jd[:MAX_TEXT_CHARS]}
"""

    return f"""
{base_rules}

Create a tailored, ATS-friendly resume for the candidate using ONLY information supported by the original resume.

TAILORING RULES:
- Prioritize relevant experience and skills from the original resume.
- Naturally incorporate important terminology from the job description ONLY when the resume supports it.
- Do not add skills simply because they appear in the JD.
- Keep all employment dates, employers, titles, degrees, certifications, projects, and facts accurate.
- Do not invent achievements or numbers.
- Use strong action verbs and concise bullet points.
- Make bullets achievement-oriented when the source contains enough evidence.
- If no evidence exists for a metric, do not invent one.
- Do not use tables, columns, graphics, icons, emojis, or decorative characters.
- Use conventional ATS sections: SUMMARY, SKILLS, EXPERIENCE, PROJECTS, EDUCATION, CERTIFICATIONS when applicable.
- Do not include a photo.
- Do not include an objective unless it adds clear value.
- Preserve the candidate's contact details exactly if they are present.
- If a section is absent from the source resume, do not fabricate it.

Return ONLY valid JSON:
{{
  "tailored_resume": {{
    "name_and_contact": "",
    "summary": "",
    "skills": [],
    "experience": [
      {{
        "company": "",
        "title": "",
        "dates": "",
        "bullets": []
      }}
    ],
    "projects": [
      {{
        "name": "",
        "description": "",
        "bullets": []
      }}
    ],
    "education": [],
    "certifications": []
  }},
  "changes_made": [],
  "keywords_added_from_jd": [],
  "keywords_not_added_because_unsupported": [],
  "tailoring_notes": []
}}

ORIGINAL RESUME:
{resume[:MAX_TEXT_CHARS]}

JOB DESCRIPTION:
{jd[:MAX_TEXT_CHARS]}
"""


def call_gemini(api_key: str, prompt: str) -> Dict[str, Any]:
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    )

    raw = (response.text or "").strip()
    if not raw:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Defensive cleanup if a provider response contains accidental code fences.
        cleaned = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.I)
        cleaned = re.sub(r"\s*```$", "", cleaned)
        return json.loads(cleaned)


def get_api_key() -> str:
    return os.getenv("GEMINI_API_KEY", "").strip()


def score_color(score: int) -> str:
    if score >= 85:
        return "🟢"
    if score >= 65:
        return "🟡"
    return "🔴"


def render_analysis(data: Dict[str, Any]) -> None:
    score = int(data.get("overall_score", 0))
    level = data.get("match_level", "Unknown")

    st.subheader("ATS Match Score")
    c1, c2, c3 = st.columns([1, 1, 2])
    c1.metric("Overall", f"{score}/100")
    c2.metric("Match Level", level)
    c3.progress(max(0, min(score, 100)) / 100, text=f"{score}% estimated match")

    st.divider()

    breakdown = data.get("score_breakdown", {})
    st.subheader("Score Breakdown")
    cols = st.columns(3)
    items = [
        ("Keyword Match", "keyword_match"),
        ("Required Skills", "required_skills"),
        ("Responsibilities", "responsibilities_match"),
        ("Experience", "experience_match"),
        ("Education / Certification", "education_certification"),
        ("ATS Formatting", "ats_formatting"),
    ]
    for i, (label, key) in enumerate(items):
        item = breakdown.get(key, {})
        cols[i % 3].metric(
            label,
            f'{item.get("score", 0)}/{item.get("max_score", 0)}',
        )
        if item.get("explanation"):
            cols[i % 3].caption(item["explanation"])

    st.subheader("Keyword Analysis")
    ka = data.get("keyword_analysis", {})
    k1, k2, k3 = st.columns(3)
    with k1:
        st.markdown("### 🟢 Matching")
        for x in ka.get("matching_keywords", []):
            st.write(f"• {x}")
    with k2:
        st.markdown("### 🟡 Partial")
        for x in ka.get("partial_keywords", []):
            st.write(f"• {x}")
    with k3:
        st.markdown("### 🔴 Missing")
        for x in ka.get("missing_keywords", []):
            st.write(f"• {x}")

    st.subheader("Strengths & Gaps")
    a, b = st.columns(2)
    with a:
        st.markdown("### Strengths")
        for x in data.get("strengths", []):
            st.write(f"✓ {x}")
    with b:
        st.markdown("### Gaps")
        for x in data.get("gaps", []):
            st.write(f"• {x}")

    st.subheader("Priority Improvements")
    for rec in data.get("recommendations", []):
        priority = rec.get("priority", "MEDIUM")
        with st.expander(f"{priority}: {rec.get('problem', 'Recommendation')}"):
            st.write(f"**Why it matters:** {rec.get('why_it_matters', '')}")
            st.write(f"**Recommendation:** {rec.get('recommendation', '')}")
            if rec.get("example"):
                st.write(f"**Example:** {rec['example']}")

    bullets = data.get("bullet_improvements", [])
    if bullets:
        st.subheader("Bullet Improvements")
        for b in bullets:
            with st.expander(b.get("original", "Resume bullet")):
                st.write("**Improved:**")
                st.write(b.get("improved", ""))
                st.caption(b.get("reason", ""))

    st.subheader("Top 5 Improvements")
    for i, x in enumerate(data.get("top_5_improvements", [])[:5], 1):
        st.write(f"**{i}.** {x}")

    if data.get("final_recommendation"):
        st.info(data["final_recommendation"])


def resume_to_text(data: Dict[str, Any]) -> str:
    r = data.get("tailored_resume", {})
    sections = []

    if r.get("name_and_contact"):
        sections.append(r["name_and_contact"])

    if r.get("summary"):
        sections += ["SUMMARY", r["summary"]]

    if r.get("skills"):
        sections += ["SKILLS", ", ".join(r["skills"])]

    if r.get("experience"):
        sections.append("EXPERIENCE")
        for exp in r["experience"]:
            heading = " | ".join(
                x for x in [exp.get("title"), exp.get("company"), exp.get("dates")] if x
            )
            sections.append(heading)
            sections.extend(f"• {b}" for b in exp.get("bullets", []))

    if r.get("projects"):
        sections.append("PROJECTS")
        for p in r["projects"]:
            sections.append(p.get("name", ""))
            if p.get("description"):
                sections.append(p["description"])
            sections.extend(f"• {b}" for b in p.get("bullets", []))

    if r.get("education"):
        sections.append("EDUCATION")
        sections.extend(f"• {x}" for x in r["education"])

    if r.get("certifications"):
        sections.append("CERTIFICATIONS")
        sections.extend(f"• {x}" for x in r["certifications"])

    return "\n\n".join(x for x in sections if x)


def render_tailored(data: Dict[str, Any]) -> None:
    st.subheader("✨ Tailored Resume")
    text = resume_to_text(data)
    if not text:
        st.error("Gemini did not return a usable tailored resume.")
        return

    st.text_area("ATS-friendly resume", text, height=650)
    st.download_button(
        "⬇️ Download Tailored Resume (.txt)",
        data=text.encode("utf-8"),
        file_name="tailored_resume.txt",
        mime="text/plain",
    )

    st.subheader("What Changed?")
    for x in data.get("changes_made", []):
        st.write(f"• {x}")

    unsupported = data.get("keywords_not_added_because_unsupported", [])
    if unsupported:
        st.subheader("JD Keywords Not Added")
        st.caption(
            "These were intentionally excluded because the original resume did not "
            "provide enough evidence to claim them."
        )
        for x in unsupported:
            st.write(f"• {x}")


def main():
    st.title("📄 ResumeMatch AI")
    st.caption("AI-powered ATS analysis and job-specific resume tailoring")

    with st.sidebar:
        st.header("Settings")
        api_key = st.text_input(
            "Gemini API Key",
            value=get_api_key(),
            type="password",
            help="You can also store it as GEMINI_API_KEY in a .env file.",
        )
        st.caption(f"Model: {MODEL}")

    left, right = st.columns(2)

    with left:
        st.subheader("1. Upload Resume")
        resume_file = st.file_uploader(
            "PDF, DOCX, or TXT",
            type=["pdf", "docx", "txt"],
            key="resume",
        )

    with right:
        st.subheader("2. Job Description")
        jd = st.text_area(
            "Paste the complete job description",
            height=280,
            placeholder="Paste the job posting here...",
        )

    if resume_file:
        try:
            resume_text = clean_text(extract_text(resume_file))
            st.success(f"Resume extracted: {len(resume_text):,} characters")
            with st.expander("Preview extracted resume text"):
                st.text(resume_text[:8000])
        except Exception as e:
            resume_text = ""
            st.error(f"Could not read the resume: {e}")
    else:
        resume_text = ""

    if len(resume_text) > MAX_TEXT_CHARS:
        st.warning(
            f"Resume is long, so only the first {MAX_TEXT_CHARS:,} characters will be sent "
            "to Gemini."
        )

    if not api_key:
        st.info("Add your Gemini API key in the sidebar or set GEMINI_API_KEY in .env.")

    if not resume_text or not jd.strip():
        st.markdown(
            "**Workflow:** Upload your resume → paste the JD → analyze → "
            "review gaps → generate a tailored resume."
        )
        return

    analyze_col, tailor_col = st.columns(2)

    with analyze_col:
        analyze_clicked = st.button(
            "🔍 Analyze ATS Score",
            use_container_width=True,
            type="primary",
        )

    with tailor_col:
        tailor_clicked = st.button(
            "✨ Generate Tailored Resume",
            use_container_width=True,
        )

    if (analyze_clicked or tailor_clicked) and not api_key:
        st.error("Please provide a Gemini API key.")
        return

    if analyze_clicked:
        with st.spinner("Analyzing your resume against the job description..."):
            try:
                result = call_gemini(
                    api_key,
                    build_prompt(resume_text, jd, "analyze"),
                )
                st.session_state["analysis"] = result
            except Exception as e:
                st.error(f"Analysis failed: {e}")

    if tailor_clicked:
        with st.spinner("Creating a truthful, job-specific tailored resume..."):
            try:
                result = call_gemini(
                    api_key,
                    build_prompt(resume_text, jd, "tailor"),
                )
                st.session_state["tailored"] = result
            except Exception as e:
                st.error(f"Tailoring failed: {e}")

    analysis = st.session_state.get("analysis")
    tailored = st.session_state.get("tailored")

    if analysis:
        st.divider()
        render_analysis(analysis)

    if tailored:
        st.divider()
        render_tailored(tailored)


if __name__ == "__main__":
    main()
